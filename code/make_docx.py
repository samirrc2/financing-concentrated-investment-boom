#!/usr/bin/env python3
"""manuscript.tex -> manuscript.docx, for journals that want a Word file.

Generated from the LaTeX source, not maintained separately, so the two cannot drift. Display
equations are rendered from the same LaTeX and embedded as images, because python-docx cannot
emit native Word equation objects. Tables are native Word tables. Run after pdflatex.
"""
import os, re, subprocess, sys, tempfile
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEX = open("manuscript.tex").read()

# ---------- display equations, rendered from the same source ----------
EQS = [
    (r"\Delta X_i \;=\; \sum_{t=2022}^{2025}\bigl(X_{it}-\bar{X}_i\bigr),", "(1)"),
    (r"\mathrm{Coverage}_i \;=\; \frac{\Delta \mathrm{OCF}_i}{\Delta \mathrm{CAPEX}_i},", "(2)"),
    (r"s_i \;=\; \frac{\Delta \mathrm{CAPEX}_i}{\sum_j \Delta \mathrm{CAPEX}_j}", ""),
    (r"\mathrm{Coverage}_G \;=\; \sum_i s_i\,\mathrm{Coverage}_i ,", "(3)"),
]

def render_equations(outdir):
    """One standalone LaTeX page per equation -> PNG at 300 dpi."""
    import pymupdf
    paths = []
    for i, (body, _) in enumerate(EQS):
        # `standalone` is not universally installed; a cropped article page is portable
        src = (r"\documentclass[12pt]{article}"
               r"\usepackage[paperwidth=10in,paperheight=2in,margin=0.15in]{geometry}"
               r"\usepackage{amsmath,amssymb,lmodern}\usepackage[T1]{fontenc}"
               r"\pagestyle{empty}\begin{document}\noindent$\displaystyle "
               + body.rstrip(",") + r"$\end{document}")
        d = tempfile.mkdtemp()
        open(f"{d}/e.tex", "w").write(src)
        subprocess.run(["pdflatex", "-interaction=nonstopmode", "-output-directory", d, f"{d}/e.tex"],
                       stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        if not os.path.exists(f"{d}/e.pdf"):
            paths.append(None); continue
        pg = pymupdf.open(f"{d}/e.pdf")[0]
        clip = pg.get_bbox() if hasattr(pg, "get_bbox") else None
        blocks = pg.get_text("blocks")
        if blocks:
            x0 = min(b[0] for b in blocks); y0 = min(b[1] for b in blocks)
            x1 = max(b[2] for b in blocks); y1 = max(b[3] for b in blocks)
            clip = pymupdf.Rect(x0 - 3, y0 - 3, x1 + 3, y1 + 3)
        out = f"{outdir}/eq{i}.png"
        pg.get_pixmap(dpi=300, clip=clip).save(out)
        paths.append(out)
    return paths

# ---------- LaTeX -> plain text ----------
REPL = [
    (r"\\ref\{tab:ladder\}", "1"), (r"\\ref\{tab:placebo\}", "2"),
    (r"\\ref\{fig:main\}", "1"), (r"\\ref\{app:supp\}", "Appendix A"),
    (r"\\eqref\{eq:cum\}", "(1)"), (r"\\eqref\{eq:cov\}", "(2)"), (r"\\eqref\{eq:agg\}", "(3)"),
    (r"\\emph\{([^}]*)\}", r"\1"), (r"\\textit\{([^}]*)\}", r"\1"), (r"\\textbf\{([^}]*)\}", r"\1"),
    (r"\\texttt\{([^}]*)\}", r"\1"), (r"\\tg\{([^}]*)\}", r"\1"),
    (r"\\url\{([^}]*)\}", r"\1"), (r"\\href\{[^}]*\}\{([^}]*)\}", r"\1"),
    (r"\\Delta", "Δ"), (r"\\%", "%"), (r"\\\$", "$"), (r"\\&", "&"), (r"\\,", " "),
    (r"\\'e", "é"), (r"``", "\u201c"), (r"''", "\u201d"),
]
def detex(s):
    s = re.sub(r"\$-(\d[\d.,]*)\$", r"−\1", s)          # $-61$ -> −61
    s = re.sub(r"\$([^$]*)\$", r"\1", s)                 # remaining inline math
    for a, b in REPL:
        s = re.sub(a, b, s)
    s = s.replace("---", "\u2014").replace("--", "\u2013")
    s = re.sub(r"\\[a-zA-Z]+\*?", "", s)
    s = s.replace("{", "").replace("}", "").replace("\\", "")
    return re.sub(r"\s+", " ", s).strip()

def paragraphs(chunk):
    for para in re.split(r"\n\s*\n", chunk):
        para = re.sub(r"(?m)^%.*$", "", para).strip()
        if not para or para.startswith("\\begin") or para.startswith("\\label"):
            continue
        yield para

def seg(a, b):
    i = TEX.index(a); return TEX[i + len(a):TEX.index(b, i)]

# ---------- build ----------
doc = Document()
st = doc.styles["Normal"]; st.font.name = "Times New Roman"; st.font.size = Pt(11)

def head(text, level):
    h = doc.add_heading(text, level=level)
    for r in h.runs: r.font.color.rgb = RGBColor(0, 0, 0)
    return h

title = re.search(r"\\title\{([^}]*)\}", TEX).group(1)
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run(detex(title)); r.bold = True; r.font.size = Pt(15)
a = doc.add_paragraph(); a.alignment = WD_ALIGN_PARAGRAPH.CENTER
a.add_run("Samir Chincholikar    Robin Chawla").font.size = Pt(11)
aff = doc.add_paragraph(); aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
aff.add_run("Independent researcher, New York, USA").italic = True

head("Abstract", 2)
doc.add_paragraph(detex(seg(r"\begin{abstract}", r"\end{abstract}")))
kw = seg(r"\begin{keyword}", r"\end{keyword}")
doc.add_paragraph("Keywords: " + detex(kw.split(r"\JEL")[0]).replace(" \\sep ", ", ").replace("sep", "\u00b7"))
doc.add_paragraph("JEL: " + detex(kw.split(r"\JEL")[1]).replace(" \\sep ", ", ").replace("sep", ","))

tmp = tempfile.mkdtemp()
eqimg = render_equations(tmp)

SECTIONS = [
    ("1. Introduction", r"\section{Introduction}", r"\section{Data and methodology}", []),
    ("2. Data and methodology", r"\section{Data and methodology}", r"\section{Results}", [0, 1, 2, 3]),
    ("3. Results", r"\section{Results}", r"\subsection{A historically", []),
]
def emit(chunk, eqs=()):
    """Paragraphs in order; equation images inserted where the LaTeX had them."""
    parts = re.split(r"(\\begin\{equation\}.*?\\end\{equation\}|\\\[.*?\\\])", chunk, flags=re.S)
    n = 0
    for part in parts:
        if part.startswith("\\begin{equation}") or part.startswith("\\["):
            if n < len(eqs) and eqimg[eqs[n]]:
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(eqimg[eqs[n]], height=Inches(0.30))
                lab = EQS[eqs[n]][1]
                if lab: p.add_run("        " + lab)
            n += 1
        else:
            for para in paragraphs(part):
                doc.add_paragraph(detex(para))

for name, a_, b_, eqs in SECTIONS:
    head(name, 1); emit(seg(a_, b_), eqs)
for name, a_, b_ in (("3.1. A historically unusual rank–coverage gradient",
                      r"\subsection{A historically unusual rank--coverage gradient}", r"\subsection{The gradient survives"),
                     ("3.2. The gradient survives sector exclusions",
                      r"\subsection{The gradient survives sector exclusions}\label{sec:comp}", r"\subsection{Limitations}"),
                     ("3.3. Limitations", r"\subsection{Limitations}\label{sec:lim}", r"\begin{table}")):
    head(name, 2); emit(seg(a_, b_))

# ---------- tables, from the LaTeX bodies so they match the PDF exactly ----------
def add_table(caption, rows, notes):
    p = doc.add_paragraph(); r = p.add_run(caption); r.bold = True
    tb = doc.add_table(rows=0, cols=len(rows[0])); tb.style = "Table Grid"
    for i, row in enumerate(rows):
        cells = tb.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = v
            if i == 0:
                for pr in cells[j].paragraphs:
                    for rr in pr.runs: rr.bold = True
    n = doc.add_paragraph(); nr = n.add_run(notes); nr.font.size = Pt(9); nr.italic = True

def latex_rows(label, ncol):
    """Data rows only: everything between \\midrule and \\bottomrule, headers and rules dropped."""
    body = TEX[TEX.index(label):]
    body = body[:body.index(r"\end{tabular}")]
    body = body[body.index(r"\midrule") + len(r"\midrule"):]
    body = body.split(r"\bottomrule")[0]
    out = []
    for line in body.split("\\\\"):
        line = re.sub(r"\\addlinespace(\[[^\]]*\])?", "", line).strip()
        if "&" not in line or "tabular" in line or "table-format" in line or "multicolumn" in line:
            continue
        cells = [detex(c).replace("-", "\u2212") if re.fullmatch(r"-?\d+", detex(c)) else detex(c)
                 for c in line.split("&")]
        if len(cells) == ncol and any(cells):
            out.append(cells)
    return out


t1 = [["", "Firms", "Share of increase (%)", "Coverage (%)"]] + latex_rows(r"\label{tab:ladder}", 4)
add_table("Table 1. Internal cash coverage of incremental investment, cumulative flows",
          t1, detex(seg(r"\textit{Notes.} Sample:", r"\end{minipage}")).replace("Notes.", "Notes. "))
t2 = [["", "2022–25", "2018–21", "2014–17"]] + latex_rows(r"\label{tab:placebo}", 4)
add_table("Table 2. Internal cash coverage by investment rank, earlier four-year windows", t2,
          detex(seg(r"\textit{Notes.} Coverage in percent", r"\end{minipage}")))

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture("figures/Figure_1.png", width=Inches(6.3))
cap = doc.add_paragraph(); cr = cap.add_run("Figure 1. " + detex(seg(r"\caption{", "}\n\\label{fig:main}")))
cr.font.size = Pt(9)

head("4. Conclusion", 1); emit(seg(r"\section{Conclusion}", r"\section*{CRediT"))
for h in ("CRediT authorship contribution statement", "Declaration of competing interest",
          "Funding", "Data availability",
          "Declaration of generative AI and AI-assisted technologies in the manuscript preparation process"):
    nxt = TEX.index("\\section*{" + h + "}") + len("\\section*{" + h + "}")
    end = min(x for x in (TEX.find("\\section*{", nxt), TEX.find("\\begin{thebibliography}", nxt)) if x > 0)
    head(h, 2); doc.add_paragraph(detex(TEX[nxt:end]))

head("References", 1)
for i, b in enumerate(re.findall(r"\\bibitem\{\w+\}\s*(.+?)(?=\n\\bibitem|\n\\end\{thebibliography\})", TEX, re.S), 1):
    doc.add_paragraph(f"[{i}] " + detex(b))

head("Appendix A. Supplementary evidence", 1)
doc.add_paragraph("The appendix tables (A1 attribution and AI-exposure robustness, A2 data "
                  "construction and validation, A3 firm-level and sectoral robustness) are set in "
                  "the typeset PDF; see manuscript.pdf.")

doc.save("manuscript.docx")
print(f"manuscript.docx written ({os.path.getsize('manuscript.docx')//1024} KB)")
